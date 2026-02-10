import hashlib
import os
import sqlite3
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from fastapi import FastAPI, Form, Request
from fastapi.responses import HTMLResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from itsdangerous import URLSafeTimedSerializer, BadSignature
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

BASE_DIR = Path(__file__).resolve().parent
STORAGE_DIR = BASE_DIR / "storage" / "generated"
STORAGE_DIR.mkdir(parents=True, exist_ok=True)
DB_PATH = BASE_DIR / "storage" / "app.db"
SECRET_KEY = os.getenv("APP_SECRET", "local-office-secret")
serializer = URLSafeTimedSerializer(SECRET_KEY)

app = FastAPI(title="Local Office Docs")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))


def init_db() -> None:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            template_type TEXT NOT NULL,
            customer_name TEXT NOT NULL,
            destination TEXT NOT NULL,
            created_at TEXT NOT NULL,
            created_by TEXT NOT NULL
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS document_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER NOT NULL,
            file_type TEXT NOT NULL,
            file_path TEXT NOT NULL,
            FOREIGN KEY(document_id) REFERENCES documents(id)
        )
        """
    )

    admin_hash = hashlib.sha256("admin123".encode()).hexdigest()
    cur.execute(
        "INSERT OR IGNORE INTO users(username, password_hash, role) VALUES (?, ?, ?)",
        ("owner", admin_hash, "owner"),
    )
    conn.commit()
    conn.close()


@app.on_event("startup")
def on_startup() -> None:
    init_db()


def current_user(request: Request):
    token = request.cookies.get("session")
    if not token:
        return None
    try:
        data = serializer.loads(token, max_age=60 * 60 * 8)
    except BadSignature:
        return None
    return data.get("username")


def create_docx(path: Path, data: dict) -> None:
    doc = Document()
    doc.add_heading("كتاب رسمي", level=1)
    doc.add_paragraph(f"الاسم: {data['customer_name']}")
    doc.add_paragraph(f"الجهة: {data['destination']}")
    doc.add_paragraph(f"التاريخ: {data['date']}")
    doc.add_paragraph(f"المرجع: {data['reference_no']}")
    doc.add_paragraph("النص الرسمي: يرجى التفضل بالاطلاع واتخاذ اللازم.")
    doc.save(path)


def create_pdf(path: Path, data: dict) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    y = 800
    lines = [
        "نموذج كتاب رسمي",
        f"الاسم: {data['customer_name']}",
        f"الجهة: {data['destination']}",
        f"التاريخ: {data['date']}",
        f"المرجع: {data['reference_no']}",
        "النص الرسمي: يرجى التفضل بالاطلاع واتخاذ اللازم.",
    ]
    for line in lines:
        c.drawString(72, y, line)
        y -= 24
    c.save()


def create_xlsx(path: Path, data: dict) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Official Letter"
    ws["A1"] = "نوع النموذج"
    ws["B1"] = "كتاب رسمي"
    ws["A2"] = "الاسم"
    ws["B2"] = data["customer_name"]
    ws["A3"] = "الجهة"
    ws["B3"] = data["destination"]
    ws["A4"] = "التاريخ"
    ws["B4"] = data["date"]
    ws["A5"] = "المرجع"
    ws["B5"] = data["reference_no"]
    wb.save(path)


@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    user = current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=303)

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, template_type, customer_name, destination, created_at FROM documents ORDER BY id DESC LIMIT 20")
    documents = cur.fetchall()
    conn.close()
    return templates.TemplateResponse("dashboard.html", {"request": request, "user": user, "documents": documents})


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request, "error": None})


@app.post("/login", response_class=HTMLResponse)
def login(request: Request, username: str = Form(...), password: str = Form(...)):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT username, password_hash FROM users WHERE username=?", (username,))
    row = cur.fetchone()
    conn.close()
    pw_hash = hashlib.sha256(password.encode()).hexdigest()
    if not row or row[1] != pw_hash:
        return templates.TemplateResponse("login.html", {"request": request, "error": "بيانات الدخول غير صحيحة"})

    token = serializer.dumps({"username": username, "exp": (datetime.utcnow() + timedelta(hours=8)).isoformat()})
    response = RedirectResponse("/", status_code=303)
    response.set_cookie("session", token, httponly=True, samesite="lax")
    return response


@app.get("/logout")
def logout():
    response = RedirectResponse("/login", status_code=303)
    response.delete_cookie("session")
    return response


@app.post("/generate")
def generate(request: Request, customer_name: str = Form(...), destination: str = Form(...), reference_no: str = Form(...)):
    user = current_user(request)
    if not user:
        return RedirectResponse("/login", status_code=303)

    now = datetime.now()
    payload = {
        "customer_name": customer_name,
        "destination": destination,
        "reference_no": reference_no,
        "date": now.strftime("%Y-%m-%d"),
    }

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO documents(template_type, customer_name, destination, created_at, created_by) VALUES (?, ?, ?, ?, ?)",
        ("official_letter", customer_name, destination, now.isoformat(), user),
    )
    doc_id = cur.lastrowid

    docx_path = STORAGE_DIR / f"doc_{doc_id}.docx"
    pdf_path = STORAGE_DIR / f"doc_{doc_id}.pdf"
    xlsx_path = STORAGE_DIR / f"doc_{doc_id}.xlsx"

    create_docx(docx_path, payload)
    create_pdf(pdf_path, payload)
    create_xlsx(xlsx_path, payload)

    for ftype, fpath in [("docx", docx_path), ("pdf", pdf_path), ("xlsx", xlsx_path)]:
        cur.execute(
            "INSERT INTO document_files(document_id, file_type, file_path) VALUES (?, ?, ?)",
            (doc_id, ftype, str(fpath.relative_to(BASE_DIR))),
        )

    conn.commit()
    conn.close()
    return RedirectResponse("/", status_code=303)


app.mount("/storage", StaticFiles(directory=str(BASE_DIR / "storage")), name="storage")
